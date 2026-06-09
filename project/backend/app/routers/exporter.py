import io
import docx
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from typing import Dict, Any, Optional
from pydantic import BaseModel

from app.routers.auth import get_current_user
from app.supabase_client import supabase_client
from app.db import UserDB
from app.schemas.json_resume_schema import (
    JSONResumeSchema, BasicsSchema, LocationSchema, ProfileSchema,
    WorkSchema, EducationSchema, SkillSchema, ProjectSchema, CertificateSchema
)
from app.schemas.ats_schema import TailoredResumeSchema
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

router = APIRouter(prefix="/exporter", tags=["Resume Exporter"])

class CombineExportRequest(BaseModel):
    tailored_resume: TailoredResumeSchema
    initial_resume: Optional[Dict[str, Any]] = None

def add_p_border_bottom(paragraph):
    """Draws a premium solid bottom border line under a paragraph in Word."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # 6/8 pt thickness
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '555555')  # Sleek medium gray
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_shading(cell, color_hex):
    """Shades the cell background with a specific hex color."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (margins) in dxa (1 dxa = 1/20 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def parse_duration(duration_str: str):
    """Helper to split a duration range (e.g. '08/2023 - 05/2026') into start and end."""
    duration_str = duration_str or ""
    start_date, end_date = "", ""
    for separator in ["–", "-", "to"]:
        if separator in duration_str:
            parts = duration_str.split(separator)
            start_date = parts[0].strip()
            end_date = parts[1].strip()
            return start_date, end_date
    return duration_str, ""

def generate_docx_resume(resume_data: dict, theme: str = "modern") -> io.BytesIO:
    """Generates a professionally-formatted Word (.docx) document styled according to the selected theme."""
    doc = docx.Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    style = doc.styles['Normal']
    font = style.font
    
    # Configure base font by theme
    if theme == "minimalist":
        font.name = 'Georgia'
    else:
        font.name = 'Calibri'
        
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    basics = resume_data.get("basics", {})
    work = resume_data.get("work", [])
    education = resume_data.get("education", [])
    skills = resume_data.get("skills", [])
    projects = resume_data.get("projects", [])
    certificates = resume_data.get("certificates", [])
    
    def add_section_header(container, title: str):
        p = container.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        if theme != "minimalist":
            add_p_border_bottom(p)
        return p

    # THEME 1: MODERN (2-Column Sidebar Layout)
    if theme == "modern":
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(2.2)
        table.columns[1].width = Inches(4.8)
        
        row = table.rows[0]
        left_cell = row.cells[0]
        right_cell = row.cells[1]
        
        set_cell_shading(left_cell, "F9FAFB")
        set_cell_margins(left_cell, top=140, bottom=140, left=180, right=180)
        set_cell_margins(right_cell, top=140, bottom=140, left=180, right=180)
        
        p_name = left_cell.add_paragraph()
        r_name = p_name.add_run(basics.get("name") or "Your Name")
        r_name.font.size = Pt(16)
        r_name.bold = True
        r_name.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        p_name.paragraph_format.space_after = Pt(2)
        
        if basics.get("label"):
            p_lbl = left_cell.add_paragraph()
            r_lbl = p_lbl.add_run(basics.get("label"))
            r_lbl.font.size = Pt(9.5)
            r_lbl.bold = True
            r_lbl.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            p_lbl.paragraph_format.space_after = Pt(8)
            
        p_contact = left_cell.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(8)
        p_contact.paragraph_format.line_spacing = 1.15
        
        contact_info = []
        if basics.get("email"): contact_info.append(basics.get("email"))
        if basics.get("phone"): contact_info.append(basics.get("phone"))
        if basics.get("url"): contact_info.append(basics.get("url"))
        loc = basics.get("location", {})
        if loc.get("city"): contact_info.append(f"{loc.get('city')}, {loc.get('region') or ''}")
        
        r_con = p_contact.add_run("\n".join(contact_info))
        r_con.font.size = Pt(8.5)
        
        profiles = basics.get("profiles", [])
        if profiles:
            p_prof = left_cell.add_paragraph()
            p_prof.paragraph_format.space_after = Pt(12)
            soc_info = []
            for prof in profiles:
                soc_info.append(f"{prof.get('network')}: {prof.get('username')}")
            r_soc = p_prof.add_run("\n".join(soc_info))
            r_soc.font.size = Pt(8.5)
            r_soc.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            
        if skills:
            add_section_header(left_cell, "SKILLS")
            for grp in skills:
                p_sk = left_cell.add_paragraph()
                p_sk.paragraph_format.space_after = Pt(4)
                r_gname = p_sk.add_run(f"{grp.get('name') or 'Skills'}:\n")
                r_gname.bold = True
                r_gname.font.size = Pt(8.5)
                r_kws = p_sk.add_run(", ".join(grp.get("keywords", [])))
                r_kws.font.size = Pt(8)
                
        if education:
            add_section_header(left_cell, "EDUCATION")
            for edu in education:
                p_edu = left_cell.add_paragraph()
                p_edu.paragraph_format.space_after = Pt(4)
                degree_text = f"{edu.get('studyType') or ''} {edu.get('area') or ''}"
                if edu.get("score"):
                    degree_text += f" ({edu.get('score')})"
                r_deg = p_edu.add_run(f"{degree_text}\n")
                r_deg.bold = True
                r_deg.font.size = Pt(8.5)
                r_inst = p_edu.add_run(f"{edu.get('institution') or ''}\n({edu.get('endDate') or ''})")
                r_inst.font.size = Pt(8)
                
        if certificates:
            add_section_header(left_cell, "CERTIFICATIONS")
            for cert in certificates:
                p_cert = left_cell.add_paragraph()
                p_cert.paragraph_format.space_after = Pt(4)
                r_cname = p_cert.add_run(f"{cert.get('name') or ''}\n")
                r_cname.bold = True
                r_cname.font.size = Pt(8.5)
                r_ciss = p_cert.add_run(f"{cert.get('issuer') or ''} {cert.get('date') or ''}")
                r_ciss.font.size = Pt(8)
                
        if basics.get("summary"):
            p_sum = right_cell.add_paragraph()
            p_sum.paragraph_format.space_after = Pt(10)
            p_sum.paragraph_format.line_spacing = 1.15
            p_sum.add_run(basics.get("summary"))
            
        if work:
            add_section_header(right_cell, "WORK EXPERIENCE")
            for job in work:
                p_job = right_cell.add_paragraph()
                p_job.paragraph_format.space_before = Pt(6)
                p_job.paragraph_format.space_after = Pt(2)
                p_job.paragraph_format.keep_with_next = True
                p_job.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                
                role = job.get("position") or "Role"
                company = job.get("name") or "Company"
                dur = f"{job.get('startDate') or ''} - {job.get('endDate') or ''}" if job.get('endDate') else (job.get('startDate') or '')
                
                r_role = p_job.add_run(f"{role} ")
                r_role.bold = True
                r_role.font.size = Pt(9.5)
                
                p_job.add_run("at ")
                r_comp = p_job.add_run(company)
                r_comp.bold = True
                r_comp.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                r_comp.font.size = Pt(9.5)
                
                p_job.add_run("\t")
                r_dur = p_job.add_run(dur)
                r_dur.italic = True
                r_dur.font.size = Pt(9)
                r_dur.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
                for hl in job.get("highlights", []):
                    p_hl = right_cell.add_paragraph(style='List Bullet')
                    p_hl.paragraph_format.space_after = Pt(2)
                    p_hl.paragraph_format.line_spacing = 1.15
                    p_hl.add_run(hl)
                    
        if projects:
            add_section_header(right_cell, "PROJECTS")
            for proj in projects:
                p_proj = right_cell.add_paragraph()
                p_proj.paragraph_format.space_before = Pt(6)
                p_proj.paragraph_format.space_after = Pt(2)
                
                pname = proj.get("name") or "Project"
                url = proj.get("url")
                
                r_name = p_proj.add_run(pname)
                r_name.bold = True
                r_name.font.size = Pt(9.5)
                
                if url:
                    p_proj.add_run(" (")
                    r_url = p_proj.add_run(url)
                    r_url.font.size = Pt(8.5)
                    r_url.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                    p_proj.add_run(")")
                    
                for hl in proj.get("highlights", []):
                    p_hl = right_cell.add_paragraph(style='List Bullet')
                    p_hl.paragraph_format.space_after = Pt(2)
                    p_hl.paragraph_format.line_spacing = 1.15
                    p_hl.add_run(hl)

    # THEME 2: PROFESSIONAL (Classic Corporate Layout with Deep Navy Accents)
    elif theme == "professional":
        # Name
        p_name = doc.add_paragraph()
        r_name = p_name.add_run(basics.get("name") or "Your Name")
        r_name.font.size = Pt(20)
        r_name.bold = True
        r_name.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)  # Slate 900
        p_name.paragraph_format.space_after = Pt(2)
        
        # Label
        if basics.get("label"):
            p_lbl = doc.add_paragraph()
            r_lbl = p_lbl.add_run(basics.get("label").upper())
            r_lbl.font.size = Pt(10)
            r_lbl.bold = True
            r_lbl.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)  # Deep Navy Blue
            p_lbl.paragraph_format.space_after = Pt(6)
            
        # Contact Row
        contact_parts = []
        if basics.get("email"): contact_parts.append(basics.get("email"))
        if basics.get("phone"): contact_parts.append(basics.get("phone"))
        if basics.get("url"): contact_parts.append(basics.get("url"))
        loc = basics.get("location", {})
        if loc.get("city"): contact_parts.append(f"{loc.get('city')}, {loc.get('region') or ''}")
        
        p_con = doc.add_paragraph()
        p_con.paragraph_format.space_after = Pt(10)
        r_con = p_con.add_run("  |  ".join(contact_parts))
        r_con.font.size = Pt(9.5)
        
        # Social profiles row
        profiles = basics.get("profiles", [])
        if profiles:
            p_prof = doc.add_paragraph()
            p_prof.paragraph_format.space_after = Pt(14)
            soc_parts = [f"{p.get('network')}: {p.get('url')}" for p in profiles if p.get('network') and p.get('url')]
            if soc_parts:
                r_prof = p_prof.add_run("  |  ".join(soc_parts))
                r_prof.font.size = Pt(9)
                r_prof.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

        def add_prof_header(title: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)  # Navy blue section header
            add_p_border_bottom(p)

        if basics.get("summary"):
            add_prof_header("PROFESSIONAL SUMMARY")
            p_sum = doc.add_paragraph()
            p_sum.paragraph_format.space_after = Pt(8)
            p_sum.paragraph_format.line_spacing = 1.15
            p_sum.add_run(basics.get("summary"))
            
        if skills:
            add_prof_header("KEY SKILLS & COMPETENCIES")
            for grp in skills:
                p_sk = doc.add_paragraph()
                p_sk.paragraph_format.space_after = Pt(3)
                p_sk.paragraph_format.line_spacing = 1.15
                r_gname = p_sk.add_run(f"{grp.get('name') or 'Skills'}: ")
                r_gname.bold = True
                p_sk.add_run(", ".join(grp.get("keywords", [])))
                
        if work:
            add_prof_header("PROFESSIONAL HISTORY")
            for job in work:
                p_job = doc.add_paragraph()
                p_job.paragraph_format.space_before = Pt(6)
                p_job.paragraph_format.space_after = Pt(2)
                p_job.paragraph_format.keep_with_next = True
                p_job.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                
                role = job.get("position") or "Role"
                company = job.get("name") or "Company"
                dur = f"{job.get('startDate') or ''} - {job.get('endDate') or ''}" if job.get('endDate') else (job.get('startDate') or '')
                
                r_role = p_job.add_run(f"{role} ")
                r_role.bold = True
                p_job.add_run("at ")
                r_comp = p_job.add_run(company)
                r_comp.bold = True
                r_comp.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                
                p_job.add_run("\t")
                r_dur = p_job.add_run(dur)
                r_dur.italic = True
                r_dur.font.size = Pt(9.5)
                r_dur.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                
                for hl in job.get("highlights", []):
                    p_hl = doc.add_paragraph(style='List Bullet')
                    p_hl.paragraph_format.space_after = Pt(2)
                    p_hl.paragraph_format.line_spacing = 1.15
                    p_hl.add_run(hl)
                    
        if projects:
            add_prof_header("KEY PROJECTS")
            for proj in projects:
                p_proj = doc.add_paragraph()
                p_proj.paragraph_format.space_before = Pt(6)
                p_proj.paragraph_format.space_after = Pt(2)
                
                pname = proj.get("name") or "Project"
                url = proj.get("url")
                
                r_name = p_proj.add_run(pname)
                r_name.bold = True
                
                if url:
                    p_proj.add_run(" (")
                    r_url = p_proj.add_run(url)
                    r_url.font.size = Pt(9)
                    r_url.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                    p_proj.add_run(")")
                    
                for hl in proj.get("highlights", []):
                    p_hl = doc.add_paragraph(style='List Bullet')
                    p_hl.paragraph_format.space_after = Pt(2)
                    p_hl.paragraph_format.line_spacing = 1.15
                    p_hl.add_run(hl)
                    
        if education:
            add_prof_header("EDUCATION")
            for edu in education:
                p_edu = doc.add_paragraph()
                p_edu.paragraph_format.space_after = Pt(3)
                p_edu.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                
                inst = edu.get("institution") or ""
                deg = edu.get("studyType") or ""
                area = edu.get("area") or ""
                year = edu.get("endDate") or ""
                score = edu.get("score")
                
                degree_str = f"{deg} in {area}" if deg and area else (deg or area)
                if score:
                    degree_str += f" ({score})"
                r_deg = p_edu.add_run(f"{degree_str} ")
                r_deg.bold = True
                p_edu.add_run("from ")
                r_inst = p_edu.add_run(inst)
                r_inst.bold = True
                
                p_edu.add_run("\t")
                r_year = p_edu.add_run(year)
                r_year.italic = True
                r_year.font.size = Pt(9.5)
                r_year.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        if certificates:
            add_prof_header("CERTIFICATIONS")
            for cert in certificates:
                p_cert = doc.add_paragraph()
                p_cert.paragraph_format.space_after = Pt(3)
                p_cert.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                
                cname = cert.get("name") or ""
                issuer = cert.get("issuer")
                date = cert.get("date")
                
                r_cert = p_cert.add_run(cname)
                r_cert.bold = True
                if issuer:
                    p_cert.add_run(" - ")
                    r_issuer = p_cert.add_run(issuer)
                    r_issuer.italic = True
                if date:
                    p_cert.add_run("\t")
                    r_date = p_cert.add_run(date)
                    r_date.italic = True
                    r_date.font.size = Pt(9.5)
                    r_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # THEME 3 & 4: ELEGANT / MINIMALIST (Standard Centered/Left Single Column)
    else:
        p_name = doc.add_paragraph()
        if theme == "elegant":
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run(basics.get("name") or "Your Name")
        r_name.font.size = Pt(20)
        r_name.bold = True
        r_name.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        
        if basics.get("label"):
            p_lbl = doc.add_paragraph()
            if theme == "elegant":
                p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_lbl = p_lbl.add_run(basics.get("label").upper())
            r_lbl.font.size = Pt(10)
            r_lbl.bold = True
            r_lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            p_lbl.paragraph_format.space_after = Pt(8)
            
        contact_parts = []
        if basics.get("email"): contact_parts.append(basics.get("email"))
        if basics.get("phone"): contact_parts.append(basics.get("phone"))
        if basics.get("url"): contact_parts.append(basics.get("url"))
        loc = basics.get("location", {})
        if loc.get("city"): contact_parts.append(f"{loc.get('city')}, {loc.get('region') or ''}")
        
        p_contact = doc.add_paragraph()
        if theme == "elegant":
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contact.paragraph_format.space_after = Pt(14)
        r_con = p_contact.add_run("  |  ".join(contact_parts))
        r_con.font.size = Pt(9.5)
        
        profiles = basics.get("profiles", [])
        if profiles:
            social_parts = [f"{p.get('network')}: {p.get('url')}" for p in profiles if p.get('network') and p.get('url')]
            if social_parts:
                p_prof = doc.add_paragraph()
                if theme == "elegant":
                    p_prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_prof.paragraph_format.space_after = Pt(18)
                r_prof = p_prof.add_run("  |  ".join(social_parts))
                r_prof.font.size = Pt(9)
                r_prof.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                
        if basics.get("summary"):
            add_section_header(doc, "SUMMARY")
            p_sum = doc.add_paragraph()
            p_sum.paragraph_format.space_after = Pt(8)
            p_sum.paragraph_format.line_spacing = 1.15
            p_sum.add_run(basics.get("summary"))
            
        if skills:
            add_section_header(doc, "TECHNICAL SKILLS")
            for grp in skills:
                p_sk = doc.add_paragraph()
                p_sk.paragraph_format.space_after = Pt(3)
                p_sk.paragraph_format.line_spacing = 1.15
                r_gname = p_sk.add_run(f"{grp.get('name') or 'Skills'}: ")
                r_gname.bold = True
                p_sk.add_run(", ".join(grp.get("keywords", [])))
                
        if work:
            add_section_header(doc, "PROFESSIONAL EXPERIENCE")
            for job in work:
                p_job = doc.add_paragraph()
                p_job.paragraph_format.space_before = Pt(6)
                p_job.paragraph_format.space_after = Pt(2)
                p_job.paragraph_format.keep_with_next = True
                p_job.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                
                role = job.get("position") or "Role"
                company = job.get("name") or "Company"
                dur = f"{job.get('startDate') or ''} - {job.get('endDate') or ''}" if job.get('endDate') else (job.get('startDate') or '')
                
                r_role = p_job.add_run(f"{role} ")
                r_role.bold = True
                p_job.add_run("at ")
                r_comp = p_job.add_run(company)
                r_comp.bold = True
                r_comp.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                
                p_job.add_run("\t")
                r_dur = p_job.add_run(dur)
                r_dur.italic = True
                r_dur.font.size = Pt(9.5)
                r_dur.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                
                for hl in job.get("highlights", []):
                    p_hl = doc.add_paragraph(style='List Bullet')
                    p_hl.paragraph_format.space_after = Pt(2)
                    p_hl.paragraph_format.line_spacing = 1.15
                    p_hl.add_run(hl)
                    
        if projects:
            add_section_header(doc, "PROJECTS")
            for proj in projects:
                p_proj = doc.add_paragraph()
                p_proj.paragraph_format.space_before = Pt(6)
                p_proj.paragraph_format.space_after = Pt(2)
                
                pname = proj.get("name") or "Project"
                url = proj.get("url")
                
                r_name = p_proj.add_run(pname)
                r_name.bold = True
                
                if url:
                    p_proj.add_run(" (")
                    r_url = p_proj.add_run(url)
                    r_url.font.size = Pt(9)
                    r_url.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                    p_proj.add_run(")")
                    
                for hl in proj.get("highlights", []):
                    p_hl = doc.add_paragraph(style='List Bullet')
                    p_hl.paragraph_format.space_after = Pt(2)
                    p_hl.paragraph_format.line_spacing = 1.15
                    p_hl.add_run(hl)
                    
        if education:
            add_section_header(doc, "EDUCATION")
            for edu in education:
                p_edu = doc.add_paragraph()
                p_edu.paragraph_format.space_after = Pt(3)
                p_edu.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                
                inst = edu.get("institution") or ""
                deg = edu.get("studyType") or ""
                area = edu.get("area") or ""
                year = edu.get("endDate") or ""
                score = edu.get("score")
                
                degree_str = f"{deg} in {area}" if deg and area else (deg or area)
                if score:
                    degree_str += f" ({score})"
                r_deg = p_edu.add_run(f"{degree_str} ")
                r_deg.bold = True
                p_edu.add_run("from ")
                r_inst = p_edu.add_run(inst)
                r_inst.bold = True
                
                p_edu.add_run("\t")
                r_year = p_edu.add_run(year)
                r_year.italic = True
                r_year.font.size = Pt(9.5)
                r_year.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        if certificates:
            add_section_header(doc, "CERTIFICATIONS")
            for cert in certificates:
                p_cert = doc.add_paragraph()
                p_cert.paragraph_format.space_after = Pt(3)
                p_cert.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                
                cname = cert.get("name") or ""
                issuer = cert.get("issuer")
                date = cert.get("date")
                
                r_cert = p_cert.add_run(cname)
                r_cert.bold = True
                if issuer:
                    p_cert.add_run(" - ")
                    r_issuer = p_cert.add_run(issuer)
                    r_issuer.italic = True
                if date:
                    p_cert.add_run("\t")
                    r_date = p_cert.add_run(date)
                    r_date.italic = True
                    r_date.font.size = Pt(9.5)
                    r_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

@router.post("/combine-and-export", response_model=JSONResumeSchema)
async def combine_and_export(
    req: CombineExportRequest,
    current_user: dict = Depends(get_current_user)
):
    user_email = current_user["email"]
    
    initial = req.initial_resume
    if not initial:
        try:
            profile_res = supabase_client.table("profiles").select("*").eq("email", user_email.lower()).execute()
            if not profile_res.data:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="No user profile found. Please upload a resume first."
                )
            profile = profile_res.data[0]
            profile_id = profile["id"]
            
            exp_res = supabase_client.table("experience").select("*").eq("profile_id", profile_id).execute()
            edu_res = supabase_client.table("education").select("*").eq("profile_id", profile_id).execute()
            proj_res = supabase_client.table("projects").select("*").eq("profile_id", profile_id).execute()
            cert_res = supabase_client.table("certifications").select("*").eq("profile_id", profile_id).execute()
            skills_res = supabase_client.table("skills").select("name").eq("profile_id", profile_id).execute()
            
            skills_list = [s["name"] for s in skills_res.data] if skills_res.data else []
            
            initial = {
                "full_name": profile.get("name", ""),
                "email": profile.get("email", ""),
                "phone": profile.get("phone", ""),
                "location": profile.get("location", ""),
                "linkedin_url": profile.get("linkedin_url", ""),
                "github_url": profile.get("github_url", ""),
                "portfolio_url": profile.get("portfolio_url", ""),
                "title": profile.get("title", ""),
                "skills": skills_list,
                "experience": exp_res.data or [],
                "projects": proj_res.data or [],
                "education": edu_res.data or [],
                "certifications": cert_res.data or []
            }
        except HTTPException as he:
            raise he
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to retrieve resume from database: {str(e)}"
            )

    tailored = req.tailored_resume

    loc_str = initial.get("location") or ""
    city, region, country = "", "", ""
    if loc_str:
        parts = [p.strip() for p in loc_str.split(",")]
        if len(parts) == 1:
            city = parts[0]
        elif len(parts) == 2:
            city = parts[0]
            region = parts[1]
        elif len(parts) >= 3:
            city = parts[0]
            region = parts[1]
            country = parts[2]

    location_schema = LocationSchema(city=city or None, region=region or None, countryCode=country or None)

    profiles_schema = []
    linkedin = initial.get("linkedin_url")
    if linkedin:
        url = linkedin if linkedin.startswith("http") else f"https://{linkedin}"
        username = linkedin.rstrip("/").split("/")[-1]
        profiles_schema.append(ProfileSchema(network="LinkedIn", username=username, url=url))
        
    github = initial.get("github_url")
    if github:
        url = github if github.startswith("http") else f"https://{github}"
        username = github.rstrip("/").split("/")[-1]
        profiles_schema.append(ProfileSchema(network="GitHub", username=username, url=url))

    basics = BasicsSchema(
        name=initial.get("full_name") or initial.get("name"),
        label=initial.get("title") or (initial.get("preferred_job_roles", [""])[0] if initial.get("preferred_job_roles") else None),
        email=initial.get("email") or user_email,
        phone=initial.get("phone"),
        url=initial.get("portfolio_url"),
        summary=tailored.professional_summary,
        location=location_schema,
        profiles=profiles_schema
    )

    work_list = []
    for exp in tailored.experience:
        start_date, end_date = parse_duration(exp.duration)
        desc = exp.description
        if isinstance(desc, list):
            highlights = desc
        elif isinstance(desc, str):
            highlights = [p.strip().lstrip("-*•").strip() for p in desc.split("\n") if p.strip()]
        else:
            highlights = []
            
        work_list.append(WorkSchema(
            name=exp.company,
            position=exp.role,
            startDate=start_date or exp.duration,
            endDate=end_date or None,
            highlights=highlights
        ))

    skills_list = []
    target_skills = tailored.skills if tailored.skills else initial.get("skills", [])
    if target_skills:
        skills_list.append(SkillSchema(
            name="Skills",
            keywords=target_skills
        ))

    project_list = []
    for proj in tailored.projects:
        desc = proj.description
        if isinstance(desc, list):
            highlights = desc
        elif isinstance(desc, str):
            highlights = [p.strip().lstrip("-*•").strip() for p in desc.split("\n") if p.strip()]
        else:
            highlights = []
            
        project_list.append(ProjectSchema(
            name=proj.name,
            url=proj.url,
            highlights=highlights,
            description=", ".join(highlights)
        ))

    edu_list = []
    for edu in tailored.education:
        start_date, end_date = parse_duration(edu.year)
        edu_list.append(EducationSchema(
            institution=edu.institution,
            studyType=edu.degree,
            startDate=start_date or None,
            endDate=end_date or edu.year,
            score=edu.gpa or None
        ))

    cert_list = []
    for cert in tailored.certifications:
        cert_list.append(CertificateSchema(
            name=cert.name,
            issuer=cert.issuer,
            date=cert.date
        ))

    return JSONResumeSchema(
        basics=basics,
        work=work_list,
        skills=skills_list,
        projects=project_list,
        education=edu_list,
        certificates=cert_list
    )

@router.post("/download/docx")
async def download_docx(
    resume: JSONResumeSchema,
    theme: str = "modern",
    current_user: dict = Depends(get_current_user)
):
    try:
        docx_bytes = generate_docx_resume(resume.model_dump(), theme)
        
        name_slug = "Resume"
        if resume.basics.name:
            name_slug = resume.basics.name.replace(" ", "_").lower()
            
        headers = {
            'Content-Disposition': f'attachment; filename="{name_slug}_{theme}_tailored.docx"'
        }
        
        return StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate Word document: {str(e)}"
        )
